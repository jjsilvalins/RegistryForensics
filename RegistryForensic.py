from winreg import *
import string
import os
import datetime
try:
    from docx import Document
    from docx.shared import Inches
    from docx.shared import Pt
except ImportError:
    print("Instale o módulo python-docx")
    print("Comando: pip install python-docx")

roots_hives = {
    "HKEY_CLASSES_ROOT": HKEY_CLASSES_ROOT,
    "HKEY_CURRENT_USER": HKEY_CURRENT_USER,
    "HKEY_LOCAL_MACHINE": HKEY_LOCAL_MACHINE,
    "HKEY_USERS": HKEY_USERS,
    "HKEY_PERFORMANCE_DATA": HKEY_PERFORMANCE_DATA,
    "HKEY_CURRENT_CONFIG": HKEY_CURRENT_CONFIG,
    "HKEY_DYN_DATA": HKEY_DYN_DATA
}

def parse_key(key):
    key = key.upper()
    parts = key.split('\\')
    root_hive_name = parts[0]
    root_hive = roots_hives.get(root_hive_name)
    partial_key = '\\'.join(parts[1:])

    if not root_hive:
        raise Exception('root hive "{}" was not found'.format(root_hive_name))

    return partial_key, root_hive


def get_sub_keys(key):
    partial_key, root_hive = parse_key(key)

    with ConnectRegistry(None, root_hive) as reg:
        with OpenKey(reg, partial_key) as key_object:
            sub_keys_count, values_count, last_modified = QueryInfoKey(key_object)
            try:
                for i in range(sub_keys_count):
                    sub_key_name = EnumKey(key_object, i)
                    yield sub_key_name
            except WindowsError:
                pass


def get_values(key, fields):
    partial_key, root_hive = parse_key(key)

    with ConnectRegistry(None, root_hive) as reg:
        with OpenKey(reg, partial_key) as key_object:
            data = {}
            for field in fields:
                try:
                    value, type = QueryValueEx(key_object, field)
                    data[field] = value
                except WindowsError:
                    pass

            return data


def get_value(key, field):
    values = get_values(key, [field])
    return values.get(field)

def join(path, *paths):
    path = path.strip('/\\')
    paths = map(lambda x: x.strip('/\\'), paths)
    paths = list(paths)
    result = os.path.join(path, *paths)
    result = result.replace('/', '\\')
    return result

def decode_binary(value):
    return value[::2][:value[::2].find(b'\x00')].decode()

INFOWINDOWS_KEY = "HKEY_LOCAL_MACHINE\SOFTWARE\Microsoft\Windows NT\CurrentVersion"
LATESTLOGGED_KEY = "HKEY_LOCAL_MACHINE\SOFTWARE\Microsoft\Windows\CurrentVersion\Authentication\LogonUI"
PROGRAMS_KEY = "HKEY_LOCAL_MACHINE\\SOFTWARE\\Microsoft\\Windows\\CurrentVersion\\Uninstall"
USB_KEY = "HKEY_LOCAL_MACHINE\\SYSTEM\\ControlSet001\\Enum\\USBSTOR"
RECENTFILES_KEY = "HKEY_CURRENT_USER\\SOFTWARE\\Microsoft\\Windows\\CurrentVersion\\Explorer\\RecentDocs"
EXECCOMMAND_KEY = "HKEY_CURRENT_USER\\SOFTWARE\\Microsoft\\Windows\\CurrentVersion\\Explorer\\RunMRU"
INFOBIOS_KEY = "HKEY_LOCAL_MACHINE\\HARDWARE\\DESCRIPTION\\System\\BIOS"
NETWORKINTEFACES_KEY = "HKEY_LOCAL_MACHINE\\SYSTEM\\CurrentControlSet\\Services\\Tcpip\\Parameters\\"
CPU_KEY = "HKEY_LOCAL_MACHINE\\HARDWARE\\DESCRIPTION\\System\\CentralProcessor\\0"

INFO_WINDOWS = get_values(INFOWINDOWS_KEY, ["ProductName","RegisteredOwner","EditionID"])
LATEST_LOGGED = get_value(LATESTLOGGED_KEY, "LastLoggedOnUser")
CPU = get_value(CPU_KEY, "ProcessorNameString")
INFO_BIOS = get_values(INFOBIOS_KEY, ["BIOSReleaseDate","BaseBoardManufacturer","BIOSVendor","BIOSVersion","SystemProductName"])

INSTALLED_PROGRAMS = []
for key in get_sub_keys(PROGRAMS_KEY):
    INSTALLED_PROGRAMS.append(get_value(PROGRAMS_KEY+"\\"+key,"DisplayName"))

DEVICES = []
for pendrives in get_sub_keys(USB_KEY):
    for pendrive in get_sub_keys(USB_KEY+"\\"+pendrives):
        DEVICES.append(get_value(USB_KEY+"\\"+pendrives+"\\"+pendrive,"FriendlyName"))

RECENT_FILES = []
for num in range(1024):
    try:
        value = get_value(RECENTFILES_KEY,str(num))
        RECENT_FILES.append(decode_binary(value))
    except:
        pass

EXEC_COMMANDS = []
for letter in string.ascii_lowercase:
    try:
        value = get_value(EXECCOMMAND_KEY,letter)
        EXEC_COMMANDS.append(value.replace("\\1",""))
    except EnvironmentError:                                
        break

#NoUserLogged, LocalSystem, LocalService, NetworkService
DEFAULT_USERS = [".DEFAULT","S-1-5-18","S-1-5-19","S-1-5-20"]
USERS = []                            
for user in get_sub_keys("HKEY_USERS"):
    if not user in DEFAULT_USERS and not "_Classes" in user: #Default users and Current User Classes
        USERS.append(get_values("HKEY_USERS\\"+user+"\\Volatile Environment",["HOMEPATH","USERNAME","USERPROFILE","USERDOMAIN"]))

NETWORK_SETTINGS = []
for adapter in get_sub_keys(NETWORKINTEFACES_KEY+"\\Adapters"):
    settings = get_values(NETWORKINTEFACES_KEY+"\\"+"Interfaces"+'\\'+adapter,["DefaultGateway", "IPAddress", "NameServer", "SubnetMask"])
    if len(settings) > 1:
        NETWORK_SETTINGS.append(settings)

document = Document()
style = document.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(12)

document.add_heading('Relatório: Registro do Windows', 0)

document.add_heading('Informações do computador')
for k, v in INFO_WINDOWS.items():
    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'{k}: ')
    run.bold = True
    paragraph.add_run(v)
    paragraph.style = 'List Bullet'
    
paragraph = document.add_paragraph()
run = paragraph.add_run(f'Último login: ')
paragraph.add_run(LATEST_LOGGED)
run.bold = True
paragraph.style = 'List Bullet'

paragraph = document.add_paragraph()
run = paragraph.add_run(f'CPU: ')
paragraph.add_run(CPU)
run.bold = True
paragraph.style = 'List Bullet'

for k, v in INFO_BIOS.items():
    paragraph = document.add_paragraph()
    run = paragraph.add_run(f'{k}: ')
    run.bold = True
    paragraph.add_run(v)
    paragraph.style = 'List Bullet'
    
document.add_heading('Usuários', level=1)
for USER in USERS:
    for k, v in USER.items():
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f'{k}: ')
        run.bold = True
        paragraph.add_run(v)
        paragraph.style = 'List Bullet'
    
document.add_heading('Arquivos acessados recentemente', level=1)
for FILE in RECENT_FILES:
    paragraph = document.add_paragraph(FILE)
    paragraph.style = 'List Bullet'

document.add_heading('Comandos no executar', level=1)
for COMMAND in EXEC_COMMANDS:
    paragraph = document.add_paragraph(COMMAND)
    paragraph.style = 'List Bullet'

INFO_BIOS = get_values(INFOBIOS_KEY, ["BIOSReleaseDate","BaseBoardManufacturer","BIOSVendor","BIOSVersion","SystemProductName"])

document.add_heading('Dispositivos USB', level=1)
for DEVICE in DEVICES:
    paragraph = document.add_paragraph(DEVICE)
    paragraph.style = 'List Bullet'
    
document.add_heading('Informações de rede', level=1)

for ADAPTER in NETWORK_SETTINGS:
    for k, v in ADAPTER.items():
        paragraph = document.add_paragraph()
        run = paragraph.add_run(f'{k}: ')
        run.bold = True
        paragraph.add_run(v)
        paragraph.style = 'List Bullet'
        
document.add_heading('Programas instalados', level=1)
for PROGRAM in INSTALLED_PROGRAMS:
    if PROGRAM != None:
        paragraph = document.add_paragraph(PROGRAM)
        paragraph.style = 'List Bullet'
    
document.save('RelatórioForense.docx')

